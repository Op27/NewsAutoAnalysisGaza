import requests # For making HTTP requests to web pages.
from bs4 import BeautifulSoup # For parsing HTML and XML documents.
import re # For regular expression operations.
from wordcloud import WordCloud # For generating word cloud images.
import matplotlib.pyplot as plt # For creating static, interactive, and animated visualizations.
from collections import Counter  # For counting hashable objects.
import docx  # For creating and updating Microsoft Word (.docx) files.
import datetime # For handling date and time operations.
from docx.shared import Pt # For specifying font size in points.
from docx.enum.text import WD_ALIGN_PARAGRAPH # For paragraph alignment options.
import os # For interacting with the operating system, like opening files.
from docx.shared import Inches # For specifying dimensions in inches, used when adding images.
from openai import OpenAI # For interacting with the OpenAI API, specifically GPT models.

# Begin script execution: scraping web data, generating visualizations, and compiling a summary into a Word document.

# Step 1: Scrape Information from the Web and save them to a text file.
print(" 📰  Collecting information from news site...")

def scrape_and_save():
    URL = 'https://www.bbc.com/news/topics/c2vdnvdg6xxt'
    response = requests.get(URL)
    soup = BeautifulSoup(response.content, 'html.parser')
    articles = soup.find_all('div', class_='sc-4befc967-0 bUlrPj')
    headlines_urls_and_descriptions = []

    # Extract headline, URL, and description for each article and save.
    for article in articles:
        headline = article.find('h2')
        description_tag = article.find('p', {'data-testid': 'card-description'})
        link = article.find('a')
        if headline and description_tag and link:
            headline_text = headline.text.strip()
            description_text = description_tag.text.strip()
            url = 'https://www.bbc.com' + link['href'] if link['href'].startswith('/') else link['href']
            headlines_urls_and_descriptions.append((headline_text, url, description_text))

    with open('articles.txt', 'w', encoding='utf-8') as file:
        for headline, url, description in headlines_urls_and_descriptions:
            file.write(f"{headline} (BBC): {description} URL:{url}\n\n")

# Step 2: Generate Visualizations (Word Cloud and Bar Chart) from the saved articles.
print(" 🎨  Creating a WordCloud and Bar Chart from collected articles...")

def generate_visualizations(file_path='articles.txt'):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    cleaned_text = re.sub(r'URL:https?\:\/\/\S+', '', text).replace('(BBC)', '')   

    # Generate and save a WordCloud image.
    wordcloud = WordCloud(width=800, height=400, background_color ='white').generate(cleaned_text)
    wordcloud.to_file('wordcloud_image.png')
    
    # Perform frequency analysis on the cleaned text and generate a bar chart.
    stop_words = set(['the', 'in', 'to', 'and', 'of', 'a', 'is', 'it', 'has', 'into', 'for', 'with', 'on', 'as', 'that', 'are', 'by', 'this', 'be', 'from', 's', 'at', 'more', 'how', 'what', 'when', 'who', 'why','was','were','but','an'])
    words = re.findall(r'\b\w+\b', cleaned_text.lower())
    filtered_words = [word for word in words if word not in stop_words]
    word_freq = Counter(filtered_words)
    top_words = word_freq.most_common(20)
    words, frequencies = zip(*top_words)
    
    plt.figure(figsize=(14, 10))
    bars = plt.bar(words, frequencies, color='skyblue')
    plt.xlabel('Word')
    plt.ylabel('Frequency')
    plt.title('Word Frequency in Text with Annotations')
    plt.xticks(rotation=45)
    for bar in bars:
        yval = bar.get_height()
        plt.text(bar.get_x() + bar.get_width() / 2, yval + 1, f'{yval}', ha='center', va='bottom')
    plt.tight_layout()
    plt.savefig('bar_chart.png')
    plt.close()  # Close the plot to prevent display in notebooks or scripts

# Sequence to execute defined functions for scraping, visualization, and summary generation.
scrape_and_save()
generate_visualizations()

# Function to generate a summary of the articles using OpenAI's GPT model.
print(" 📝  Summarizing the collected articles...")

OPENAI_API_KEY = "ADD_YOUR_API_KEY_HERE"  # Replace with your actual OpenAI API key

def generate_text_with_gpt(prompt_path):
    client = OpenAI(api_key=OPENAI_API_KEY)
    with open(prompt_path, 'r', encoding='utf-8') as file:
        custom_prompt = file.read()
    with open('articles.txt', 'r', encoding='utf-8') as file:
        articles_text = file.read()
    
    # Use OpenAI's GPT to generate a summary based on the articles' text.
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",  # Adjust based on available models
        messages=[
            {"role": "system", "content": custom_prompt},
            {"role": "user", "content": articles_text}
        ]
    )
    return completion.choices[0].message.content if completion.choices else "No summary available."

prompt_result = generate_text_with_gpt('prompt.txt')

# Main function to compile the generated outputs into a Word document.
def compile_outputs_into_word(prompt_result, wordcloud_path='wordcloud_image.png', bar_chart_path='bar_chart.png', articles_path='articles.txt'):
    doc = docx.Document()

    # Setting up document style and font.
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Roboto'
    font.size = Pt(11) 

    # Include today's date in the title.
    today_date = datetime.datetime.now().strftime('%Y/%m/%d')
    title = "Summary of BBC articles on Gaza " + today_date
    title_paragraph = doc.add_heading(level=0)
    title_run = title_paragraph.add_run(title)  

    title_run.font.name = 'Roboto'
    title_run.font.size = Pt(24)  # Adjust the size as per your needs
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add wordcloud image
    doc.add_paragraph("Table 1: Frequency of words used in articles with WordCloud")
    doc.add_picture(wordcloud_path, width=Inches(5.6))
    
    # Add bar chart
    doc.add_paragraph("Table 2: Frequency of words used in articles with Bar chart")
    doc.add_picture(bar_chart_path, width=Inches(5.6))
    
    # Add a section for the generated summary text.
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Summary of articles:")
    run.bold = True
    
    # List all the scraped articles at the end of the document.
    doc.add_paragraph(prompt_result)
    
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("List of articles:")
    run.bold = True

    with open(articles_path, 'r', encoding='utf-8') as file:
        articles_text = file.read()
    doc.add_paragraph(articles_text)

    # Save the compiled document with a dynamic filename.
    file_name = 'Article summary_' + today_date.replace("/", "") + ".docx"
    doc.save(file_name)
    return file_name  # Return the file name

# Execute the compilation function and automatically open the generated Word document.
file_name = compile_outputs_into_word(prompt_result)

# Open the Word document automatically
os.startfile(file_name) # Note: os.startfile() works only on Windows.

print(" ✅  All processes completed. Outputs compiled.")